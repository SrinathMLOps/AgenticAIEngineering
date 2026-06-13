"""
generate_docs.py
─────────────────────────────────────────────────────
Generates recap_and_resources.docx — a structured Word
document for teaching Transformer Internals + Prompt Anatomy.

Run: python generate_docs.py
Prereq: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

# ─── Colour Palette ──────────────────────────────────────────────
BLUE_DARK   = RGBColor(15,  23,  42)   # slate-900
BLUE_MID    = RGBColor(2,  132, 199)   # sky-600
BLUE_LIGHT  = RGBColor(186, 230, 253)  # sky-200
ACCENT      = RGBColor(99,  102, 241)  # indigo-500
GREEN       = RGBColor(21,  128,  61)  # green-700
AMBER       = RGBColor(180,  83,   9)  # amber-800
GRAY        = RGBColor(100, 116, 139)  # slate-500
WHITE       = RGBColor(255, 255, 255)

def shade_cell(cell, hex_color: str):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("w:top",top),("w:bottom",bottom),("w:left",left),("w:right",right)]:
        node = OxmlElement(side)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.bold  = True
    r.font.color.rgb = BLUE_MID if level == 1 else ACCENT
    r.font.size  = Pt(16) if level == 1 else Pt(13)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def body(doc, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.color.rgb = BLUE_DARK
    r2 = p.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.color.rgb = BLUE_MID
    r2 = p.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def code_block(doc, code: str):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = AMBER
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shd)
    return p

def divider(doc):
    p = doc.add_paragraph("─" * 70)
    p.runs[0].font.color.rgb = BLUE_LIGHT
    p.runs[0].font.size = Pt(8)
    p.paragraph_format.space_after = Pt(2)

# ─────────────────────────────────────────────────────────────────
# DOCUMENT BUILD
# ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Page margins ──
    for sec in doc.sections:
        sec.top_margin    = Inches(0.9)
        sec.bottom_margin = Inches(0.9)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    # ═══════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Transformer Internals + Prompt Anatomy")
    r.font.name = "Calibri"; r.font.size = Pt(26); r.font.bold = True
    r.font.color.rgb = BLUE_MID

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Teaching Guide — Agentic AI Focus")
    r2.font.name = "Calibri"; r2.font.size = Pt(13); r2.font.italic = True
    r2.font.color.rgb = GRAY

    doc.add_paragraph()

    # ── Lesson overview table ──
    table = doc.add_table(rows=9, cols=3)
    table.style = "Table Grid"
    headers = ["Level", "Lesson", "Focus"]
    rows_data = [
        ("Beginner",     "L1: Build Attention From Scratch",      "NumPy → Q,K,V → real API comparison"),
        ("Beginner",     "L2: 5-Variant Prompt Harness",          "Live evaluation + scoring leaderboard"),
        ("Intermediate", "L3: Agentic Tool-Calling Loop",         "Native Gemini function calling + dispatch"),
        ("Intermediate", "L4: Position Ablation Study",           "Lost-in-the-Middle experiment + metrics table"),
        ("Intermediate", "L7: JSON Prompting Mastery",            "5 techniques, schema validation, auto-repair loop"),
        ("Advanced",     "L5: Multi-Agent Pipeline",              "Planner → Executor → Critic (3 agents)"),
        ("Advanced",     "L6: Attention Visualisation",           "GPT-2 attention heatmaps, entropy metrics"),
        ("All",          "Coding Challenges",                     "3 extensions per lesson"),
    ]

    # Header row
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        shade_cell(c, "0284C7")
        set_cell_margins(c)
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = "Calibri"

    # Data rows
    level_colors = {"Beginner":"E0F2FE","Intermediate":"EDE9FE","Advanced":"DCFCE7","All":"F1F5F9"}
    for row_idx, (level, lesson, focus) in enumerate(rows_data, 1):
        row = table.rows[row_idx]
        bg  = level_colors.get(level, "FFFFFF")
        for ci, text in enumerate([level, lesson, focus]):
            c = row.cells[ci]
            shade_cell(c, bg)
            set_cell_margins(c)
            r = c.paragraphs[0].add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(10)
            if ci == 0: r.font.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 1 — TRANSFORMER INTERNALS CHEAT SHEET
    # ═══════════════════════════════════════════════════════════
    heading(doc, "1. Transformer Internals — Core Concepts")
    divider(doc)

    heading(doc, "1.1  Scaled Dot-Product Attention", level=2)
    body(doc, "The fundamental computation in every transformer layer:")
    code_block(doc, "Attention(Q, K, V)  =  softmax( Q·Kᵀ / √dₖ )  ·  V")
    body(doc, "")

    # QKV table
    qkv_table = doc.add_table(rows=4, cols=2)
    qkv_table.style = "Table Grid"
    qkv_data = [
        ("Symbol", "Role & Prompt Engineering Implication"),
        ("Q  (Query)",  "What the model is 'searching for'. When generating a JSON key, Q searches back for relevant schema tokens."),
        ("K  (Key)",    "The 'index' of each token. Well-structured delimiters (< >, { }) produce distinctive K vectors that stand out."),
        ("V  (Value)",  "The actual content copied to the output. Few-shot examples provide rich V vectors for induction heads to copy."),
    ]
    for ri, (a, b) in enumerate(qkv_data):
        row = qkv_table.rows[ri]
        if ri == 0:
            for ci, text in enumerate([a, b]):
                shade_cell(row.cells[ci], "1E293B")
                set_cell_margins(row.cells[ci])
                r = row.cells[ci].paragraphs[0].add_run(text)
                r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = "Calibri"
        else:
            for ci, text in enumerate([a, b]):
                shade_cell(row.cells[ci], "F8FAFC" if ri%2==1 else "F1F5F9")
                set_cell_margins(row.cells[ci])
                r = row.cells[ci].paragraphs[0].add_run(text)
                r.font.name = "Calibri"; r.font.size = Pt(10)
                if ci == 0: r.font.bold = True; r.font.color.rgb = BLUE_MID

    doc.add_paragraph()

    heading(doc, "1.2  Multi-Head Attention", level=2)
    body(doc, ("Runs attention h times in parallel over sub-spaces of d_model. "
               "Each head captures a different relationship type:"))
    for h_txt in ["Head 0–2: syntactic structure (subject, verb, object)",
                  "Head 3–5: semantic similarity and coreference",
                  "Head 6–9: positional / format patterns (JSON structure, XML tags)",
                  "Head 10+: long-range dependency (tool name ↔ parameter slot)"]:
        bullet(doc, h_txt)

    code_block(doc, "MultiHead(Q,K,V) = Concat(head₁, …, headₕ) · W_O\nhead_i = Attention(Q·W_Qᵢ, K·W_Kᵢ, V·W_Vᵢ)")

    heading(doc, "1.3  Induction Heads (Why Few-Shot Works)", level=2)
    body(doc, ("Induction heads are attention heads trained to find the pattern "
               "[A][B]…[A] → predict [B]. When you give few-shot examples:"))
    code_block(doc,
               'User: Weather in London?\nOutput: {"tool":"get_weather","location":"London"}\n'
               'User: Weather in Paris?        # model sees [London?]→[Output:] pattern\n'
               'Output: {"tool":"get_weather","location":"Paris"}  # induction head copies it')

    heading(doc, "1.4  Positional Encoding", level=2)
    body(doc, ("Without position information, attention treats tokens as a bag-of-words. "
               "Positional encodings inject order. Implication for prompting:"))
    bullet(doc, "Tokens at position 0 and position N-1 receive the highest positional signal.")
    bullet(doc, "This creates U-shaped recall (primacy + recency bias).")
    bullet(doc, "Place critical constraints at the END of the system prompt.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 2 — PROMPT ANATOMY CHEAT SHEET
    # ═══════════════════════════════════════════════════════════
    heading(doc, "2. Prompt Anatomy — 5 Structural Patterns")
    divider(doc)

    patterns = [
        ("A. Barebones",
         "\"You are a tool-calling agent. Call get_weather when asked.\"",
         "Low — no format anchor, model defaults to prose",
         "Fast to write, quick demos only"),
        ("B. Few-Shot",
         "User: weather London? → {\"tool\":\"get_weather\",\"location\":\"London\"}\nUser: weather Paris? → ...",
         "High — induction heads copy the example format exactly",
         "Best for reliable JSON/structured output"),
        ("C. Chain-of-Thought",
         "Step 1: identify city. Step 2: identify unit. Step 3: output JSON only.",
         "Medium-High — scratchpad reduces hallucination",
         "Complex multi-step reasoning tasks"),
        ("D. XML-Delimited",
         "<tool>get_weather</tool><param>city</param>",
         "High — delimiters act as attention sinks for format heads",
         "Structural APIs, agent frameworks"),
        ("E. Diluted",
         "Policy 1…Policy 2…[buried rule]…Policy 7…",
         "Low — attention budget split across all policy tokens",
         "Anti-pattern to teach / diagnose"),
    ]

    pt = doc.add_table(rows=len(patterns)+1, cols=4)
    pt.style = "Table Grid"
    for ci, h_txt in enumerate(["Pattern","Example","Attention Effect","Best Use"]):
        c = pt.rows[0].cells[ci]
        shade_cell(c, "4338CA")
        set_cell_margins(c)
        r = c.paragraphs[0].add_run(h_txt)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = "Calibri"; r.font.size = Pt(10)

    pattern_colors = ["E0F2FE","EDE9FE","DCFCE7","FEF9C3","FEE2E2"]
    for ri, (name, ex, effect, use) in enumerate(patterns):
        row = pt.rows[ri+1]
        bg  = pattern_colors[ri]
        for ci, text in enumerate([name, ex, effect, use]):
            c = row.cells[ci]
            shade_cell(c, bg)
            set_cell_margins(c)
            r = c.paragraphs[0].add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(9)
            if ci == 0: r.font.bold = True

    doc.add_paragraph()

    heading(doc, "2.1  Critical Prompt Design Rules", level=2)
    rules = [
        ("Rule 1 — End Anchoring: ",    "Place the tool schema and format rule at the VERY END of the system prompt."),
        ("Rule 2 — No Dilution: ",      "Remove any policy not directly relevant to the task. Each extra sentence steals attention budget."),
        ("Rule 3 — Use Delimiters: ",   "XML tags, triple-backtick blocks, and JSON braces create strong attention anchors."),
        ("Rule 4 — Add One Example: ",  "A single high-quality few-shot example often outperforms 5 paragraphs of instructions."),
        ("Rule 5 — CoT for Complexity: ","Use 'think step-by-step' prefix when the model must extract multiple parameters or reason."),
    ]
    for prefix, text in rules:
        bullet(doc, text, bold_prefix=prefix)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 3 — AGENTIC AI PATTERNS
    # ═══════════════════════════════════════════════════════════
    heading(doc, "3. Agentic AI — Core Loop Patterns")
    divider(doc)

    heading(doc, "3.1  Basic Tool-Call Loop", level=2)
    code_block(doc,
               "loop:\n"
               "  response = model.generate(history + tools)\n"
               "  if response.has_function_call:\n"
               "      result = execute_tool(response.function_call)\n"
               "      history.append(tool_result)\n"
               "  else:\n"
               "      return response.text   # final answer")

    heading(doc, "3.2  Multi-Agent Pipeline (Lesson 5)", level=2)
    body(doc, "Three-agent pattern — each agent has a different system prompt:")
    for agent_name, desc in [
        ("PLANNER: ", "Decomposes goal into tool call steps → outputs JSON plan"),
        ("EXECUTOR: ","Dispatches each step to Python functions → collects results"),
        ("CRITIC: ",  "Reviews trace, checks completeness, writes user-facing summary"),
    ]:
        bullet(doc, desc, bold_prefix=agent_name)

    heading(doc, "3.3  Context Window Pressure", level=2)
    body(doc, "Each tool result appended to history grows the context. Watch for:")
    bullet(doc, "Attention dilution: early task goal gets 'buried' under tool results")
    bullet(doc, "Fix: Re-state the original goal in the tool-result message")
    bullet(doc, "Fix: Summarise tool results before appending them to history")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 4 — JSON PROMPTING
    # ═══════════════════════════════════════════════════════════
    heading(doc, "4. JSON Prompting — Getting Reliable JSON from LLMs")
    divider(doc)

    body(doc,
         "In agentic AI, JSON is the universal contract between the model and your code. "
         "Every tool call, structured extraction, and agent-to-agent message depends on "
         "valid, parseable JSON. LLMs are NOT guaranteed to produce it — this section "
         "shows why, and how to fix it.")

    heading(doc, "4.1  Why JSON Breaks — 6 Common Failure Modes", level=2)
    failures = [
        ("1. Trailing commas: ",       '{"a":1,"b":2,}   — common pattern in Python training data'),
        ("2. Single quotes: ",         "{\u2018key\u2019: \u2018val\u2019}   — Python dict syntax leaking into output"),
        ("3. Prose wrapping: ",        '"Here is the JSON: {...}"  — conversational padding around output'),
        ("4. Markdown fencing: ",      '```json\n{...}\n```  — model formats for readability, not parsing'),
        ("5. Truncated output: ",      '{"key": "very long value...  — token limit hit mid-generation'),
        ("6. Mismatched braces: ",     '{"a": {"b": 1}  — missing closing brace from nested objects'),
    ]
    for prefix, text in failures:
        bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()
    heading(doc, "4.2  JSON Extraction — 5-Strategy Fallback Chain", level=2)
    body(doc, "Always attempt multiple extraction strategies in order of strictness:")
    code_block(doc,
               "def extract_json(text):\n"
               "    # 1. Strip markdown fences\n"
               "    text = re.sub(r'```json', '', text).strip()\n"
               "    # 2. Direct parse\n"
               "    try: return json.loads(text)\n"
               "    except: pass\n"
               "    # 3. Find first {...} block\n"
               "    m = re.search(r'(\\{[\\s\\S]*\\})', text)\n"
               "    if m: return json.loads(m.group(1))\n"
               "    # 4. Find first [...] block\n"
               "    m = re.search(r'(\\[[\\s\\S]*\\])', text)\n"
               "    if m: return json.loads(m.group(1))\n"
               "    # 5. Fix trailing commas\n"
               "    fixed = re.sub(r',\\s*([\\}\\]])', r'\\1', text)\n"
               "    return json.loads(fixed)")

    heading(doc, "4.3  The 5 JSON Prompting Techniques", level=2)

    json_techniques = [
        ("T1. Bare Instruction",
         "'Output JSON'",
         "Low — model defaults to prose + markdown fencing",
         "Never use alone for structured output"),
        ("T2. Schema-in-Prompt",
         "List required fields + types in the system prompt",
         "Medium — model sees the shape, reduces missing keys",
         "Simple single-object extractions"),
        ("T3. One-Shot Example",
         "Input: [text] → Output: {\"key\": \"val\", ...}",
         "High — induction heads copy exact format",
         "Best general-purpose JSON extraction"),
        ("T4. Chain-of-Thought → JSON",
         "Step 1: find X. Step 2: find Y. Step 3: output JSON",
         "High — scratchpad resolves ambiguity before output",
         "Complex multi-field extractions"),
        ("T5. Constrained Prefix",
         "Prompt ends with '{' forcing model inside JSON",
         "Very High — model cannot output prose before the JSON",
         "When absolute format compliance is needed"),
    ]

    jt = doc.add_table(rows=len(json_techniques)+1, cols=4)
    jt.style = "Table Grid"
    for ci, h_txt in enumerate(["Technique", "How It Works", "Compliance", "Best For"]):
        c = jt.rows[0].cells[ci]
        shade_cell(c, "065F46")  # emerald-800
        set_cell_margins(c)
        r = c.paragraphs[0].add_run(h_txt)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = "Calibri"; r.font.size = Pt(10)

    jt_colors = ["F0FDF4", "DCFCE7", "BBF7D0", "A7F3D0", "D1FAE5"]
    for ri, (name, how, compliance, best) in enumerate(json_techniques):
        row = jt.rows[ri+1]
        bg  = jt_colors[ri % len(jt_colors)]
        for ci, text in enumerate([name, how, compliance, best]):
            c = row.cells[ci]
            shade_cell(c, bg)
            set_cell_margins(c)
            r = c.paragraphs[0].add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(9)
            if ci == 0: r.font.bold = True

    doc.add_paragraph()

    heading(doc, "4.4  Schema Validation + Auto-Repair Loop", level=2)
    body(doc, "Never trust raw LLM output. Always validate and auto-repair:")
    code_block(doc,
               "# Step 1: Extract JSON\n"
               "data = extract_json(llm_response)\n\n"
               "# Step 2: Validate against schema\n"
               "valid, errors = validate_schema(data, required_keys, types)\n\n"
               "# Step 3: Auto-repair if needed (ask model to fix its own output)\n"
               "if not valid:\n"
               "    repair_prompt = f'Fix this JSON: {llm_response}\\nSchema: {schema}'\n"
               "    data = extract_json(model.generate(repair_prompt))\n\n"
               "# Step 4: Use validated data safely\n"
               "process(data)")

    heading(doc, "4.5  Gemini JSON Mode (Most Reliable — Production Use)", level=2)
    body(doc, "Gemini supports constrained decoding that guarantees valid JSON output:"
              " pass response_mime_type='application/json' to GenerationConfig.")
    code_block(doc,
               "model = genai.GenerativeModel(\n"
               "    'gemini-1.5-flash',\n"
               "    system_instruction='Extract hotel, city, ratings as JSON.',\n"
               "    generation_config=genai.GenerationConfig(\n"
               "        response_mime_type='application/json'\n"
               "    )\n"
               ")\n"
               "resp = model.generate_content(review_text)\n"
               "data = json.loads(resp.text)   # always valid JSON")

    body(doc, "Use this for production systems. Use T1-T5 techniques for testing, education, and non-Gemini APIs.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 5 — CURATED RESOURCES
    # ═══════════════════════════════════════════════════════════
    heading(doc, "5. Curated Resources")
    divider(doc)

    # ── Papers ────────────────────────────────────────────────
    heading(doc, "4.1  Essential Papers", level=2)
    papers = [
        ("Attention Is All You Need",
         "Vaswani et al., 2017",
         "arxiv.org/abs/1706.03762",
         "Original transformer paper. Must-read for Q, K, V, Multi-Head and Positional Encoding."),
        ("Lost in the Middle",
         "Liu et al., 2023",
         "arxiv.org/abs/2307.03172",
         "Proves U-shaped recall in long-context models. Motivates end-anchoring of constraints."),
        ("Toolformer",
         "Schick et al., 2023",
         "arxiv.org/abs/2302.04761",
         "Shows how LLMs can be trained to self-insert API calls using delimiter syntax."),
        ("Chain-of-Thought Prompting",
         "Wei et al., 2022",
         "arxiv.org/abs/2201.11903",
         "Establishes CoT as a reasoning strategy; shows scratchpad tokens improve multi-step tasks."),
        ("In-Context Learning (ICL)",
         "Brown et al., 2020 (GPT-3)",
         "arxiv.org/abs/2005.14165",
         "Introduces few-shot prompting at scale; reveals induction head behaviour."),
        ("A Mathematical Framework for Transformer Circuits",
         "Elhage et al., 2021",
         "transformer-circuits.pub/2021/framework",
         "Mechanistic interpretability: explains induction heads, attention patterns, QK/OV circuits."),
    ]

    pt2 = doc.add_table(rows=len(papers)+1, cols=4)
    pt2.style = "Table Grid"
    for ci, h_txt in enumerate(["Title","Authors / Year","Link","Why It Matters"]):
        c = pt2.rows[0].cells[ci]
        shade_cell(c, "0F172A")
        set_cell_margins(c)
        r = c.paragraphs[0].add_run(h_txt)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = "Calibri"; r.font.size = Pt(10)

    paper_colors = ["F0F9FF","FEF3C7"]
    for ri, (title, authors, link, why) in enumerate(papers):
        row = pt2.rows[ri+1]
        bg  = paper_colors[ri % 2]
        for ci, text in enumerate([title, authors, link, why]):
            c = row.cells[ci]
            shade_cell(c, bg)
            set_cell_margins(c)
            r = c.paragraphs[0].add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(9)
            if ci == 0: r.font.bold = True
            if ci == 2: r.font.color.rgb = BLUE_MID

    doc.add_paragraph()

    # ── Tools & Libraries ─────────────────────────────────────
    heading(doc, "4.2  Tools & Libraries", level=2)

    tools_data = [
        ("Google AI Studio",      "aistudio.google.com",         "Live prompt playground for Gemini — test all 5 variants interactively"),
        ("Google Gen AI SDK",     "github.com/google-gemini/generative-ai-python",  "Official Python SDK used in Lessons 2–5"),
        ("TransformerLens",       "github.com/neelnanda-io/TransformerLens",   "Inspect any transformer's residual stream, attention heads and circuits"),
        ("BertViz",               "github.com/jessevig/bertviz",               "Interactive attention head visualiser for BERT/GPT family"),
        ("LangChain",             "python.langchain.com",                      "Framework for building agentic pipelines with tool calling"),
        ("LangGraph",             "langchain-ai.github.io/langgraph",          "Graph-based multi-agent orchestration (extension of Lesson 5)"),
        ("HuggingFace Transformers","huggingface.co/docs/transformers",        "Load GPT-2 / Llama for local attention extraction (Lesson 6)"),
        ("OpenAI Cookbook",       "cookbook.openai.com",                       "Practical agentic patterns, function calling recipes, RAG guides"),
    ]

    tl = doc.add_table(rows=len(tools_data)+1, cols=3)
    tl.style = "Table Grid"
    for ci, h_txt in enumerate(["Tool / Library","URL","Use Case in This Course"]):
        c = tl.rows[0].cells[ci]
        shade_cell(c, "0F172A")
        set_cell_margins(c)
        r = c.paragraphs[0].add_run(h_txt)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = "Calibri"; r.font.size = Pt(10)

    for ri, (name, url, use) in enumerate(tools_data):
        row = tl.rows[ri+1]
        bg  = "F0FFF4" if ri%2==0 else "F8FAFC"
        for ci, text in enumerate([name, url, use]):
            c = row.cells[ci]
            shade_cell(c, bg)
            set_cell_margins(c)
            r = c.paragraphs[0].add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(9)
            if ci == 0: r.font.bold = True
            if ci == 1: r.font.color.rgb = BLUE_MID

    doc.add_paragraph()

    # ── Videos & Courses ──────────────────────────────────────
    heading(doc, "4.3  Videos & Courses", level=2)

    videos = [
        ("Andrej Karpathy — Let's build GPT from scratch",
         "youtube.com/watch?v=kCc8FmEb1nY",
         "90-min video. Builds a full transformer from scratch. Best beginner video on internals."),
        ("3Blue1Brown — Attention in Transformers",
         "youtube.com/watch?v=eMlx5fFNoYc",
         "Visual walkthrough of Q, K, V and softmax. Ideal for classroom intro (15 min)."),
        ("Andrej Karpathy — Neural Nets: Zero to Hero",
         "karpathy.ai/zero-to-hero.html",
         "Full series from micrograd → GPT. Covers backprop, RNNs, attention, transformers."),
        ("DeepLearning.AI — Prompt Engineering for Devs",
         "deeplearning.ai/short-courses/chatgpt-prompt-engineering",
         "Andrew Ng + OpenAI. Covers few-shot, CoT, formatting, and tool use with examples."),
        ("Neel Nanda — Mechanistic Interpretability",
         "youtube.com/@neelnanda2469",
         "Deep dives into induction heads, attention patterns, and circuit-level analysis."),
        ("Gemini API Docs — Function Calling",
         "ai.google.dev/gemini-api/docs/function-calling",
         "Official guide used in Lessons 3–5. Native tool schema declaration + dispatch."),
    ]

    for title, url, desc in videos:
        p = doc.add_paragraph(style="List Bullet")
        r_title = p.add_run(f"{title}  ")
        r_title.font.bold = True; r_title.font.name = "Calibri"; r_title.font.size = Pt(10)
        r_url = p.add_run(f"({url})  ")
        r_url.font.color.rgb = BLUE_MID; r_url.font.name = "Calibri"; r_url.font.size = Pt(9)
        r_desc = p.add_run(f"— {desc}")
        r_desc.font.name = "Calibri"; r_desc.font.size = Pt(9); r_desc.font.color.rgb = GRAY

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 6 — RECAP QUIZ
    # ═══════════════════════════════════════════════════════════
    heading(doc, "6. Revision Quiz — 13 Questions")
    divider(doc)

    quiz = [
        ("Q1",  "What does the softmax in Attention(Q,K,V) = softmax(QKᵀ/√dₖ)·V ensure?",
                "Attention weights across all keys sum to 1.0, making them a probability distribution."),
        ("Q2",  "Why do we divide by √dₖ before applying softmax?",
                "Large dot products push softmax into saturation (near-zero gradients). Scaling keeps variance ≈ 1."),
        ("Q3",  "What are induction heads, and which prompt technique exploits them?",
                "Attention heads that complete [A][B]…[A]→[B] patterns. Few-shot examples activate them."),
        ("Q4",  "Describe the 'Lost in the Middle' phenomenon.",
                "Models have U-shaped recall: strong at start/end, weak in the middle. Critical constraints placed centrally are often ignored."),
        ("Q5",  "What does a low attention entropy value tell you?",
                "The model is focused: most attention weight is concentrated on a small number of tokens."),
        ("Q6",  "In the agentic loop, what happens after the model emits a FunctionCall?",
                "The caller executes the function, appends the FunctionResponse to history, and calls the model again."),
        ("Q7",  "Why do XML delimiters (<tool>, <call>) improve structured output compliance?",
                "They produce distinctive Key vectors that attract format-specialised attention heads, acting as anchors."),
        ("Q8",  "What is the role of the CRITIC agent in the 3-agent pipeline?",
                "It reviews the execution trace, checks that all user sub-goals were addressed, and writes the final user-facing summary."),
        ("Q9",  "What is the practical difference between d_k=1 and d_k=512 in attention?",
                "d_k=1: scores are tiny, softmax gives near-uniform weights (unfocused). d_k=512: without √dₖ scaling, scores explode and softmax becomes a near-one-hot spike."),
        ("Q10", "Name two techniques to mitigate context window pressure in long agent loops.",
                "1) Summarise tool results before appending. 2) Re-state the original goal each turn to prevent primacy decay."),
        ("Q11", "What are the 3 most common reasons LLM JSON output fails to parse?",
                "Trailing commas, prose wrapping around the JSON, and markdown fencing (``` blocks)."),
        ("Q12", "What is the 'Constrained Prefix' technique and why is it effective?",
                "The system prompt ends with '{' so the model begins generation inside the JSON — it cannot prepend prose. This exploits the model's autoregressive nature."),
        ("Q13", "When should you use response_mime_type='application/json' in Gemini vs. prompt-based techniques?",
                "Use JSON mode for production systems needing guaranteed valid JSON. Use prompt techniques (T1-T5) for learning, cross-API work, or when you need schema control in the prompt itself."),
    ]

    for q_id, question, answer in quiz:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        rq = p.add_run(f"{q_id}.  {question}")
        rq.font.bold = True; rq.font.name = "Calibri"; rq.font.size = Pt(11)
        rq.font.color.rgb = BLUE_DARK

        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Inches(0.35)
        ra = pa.add_run(f"▶  {answer}")
        ra.font.name = "Calibri"; ra.font.size = Pt(10)
        ra.font.color.rgb = GREEN

    # ── Footer ──
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Day 2 — Transformer Internals + Prompt Anatomy + JSON Prompting  |  Agentic AI Teaching Guide")
    fr.font.name = "Calibri"; fr.font.size = Pt(9); fr.font.color.rgb = GRAY

    # ── Save ──
    path = "recap_and_resources.docx"
    doc.save(path)
    print(f"[+] Word document saved: {path}")
    print(f"    Sections: Title | Transformer Internals | Prompt Anatomy | Agentic Patterns | JSON Prompting | Resources | Quiz (13 Qs)")

if __name__ == "__main__":
    build()
